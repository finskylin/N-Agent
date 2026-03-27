"""
Document Reader Skill
文档读取与识别技能 — 支持 PDF / DOCX / XLSX / 图片

依赖: httpx, pdfplumber, pypdf, python-docx, openpyxl, pytesseract, Pillow

文件来源：
  1. 本地文件路径 (file_path)
  2. HTTP URL (file_url)
  3. 钉钉附件 (dingtalk_download_code + dingtalk_robot_code)

图片识别策略：
  - 优先使用视觉大模型（VL1 → VL2 → VL3，circuit-breaker failover）
  - 所有视觉模型熔断时降级到 pytesseract OCR
  - 熔断阈值：连续失败 >= LLM_FAILOVER_THRESHOLD_VL 次，冷却 LLM_FAILOVER_COOLDOWN_VL 秒后自动恢复
"""

import asyncio
import base64
import json
import os
import sys
import tempfile
import threading
import time
from pathlib import Path
from typing import Any, Dict, List, Optional

from loguru import logger

# ========== 常量 ==========

UPLOAD_DIR = Path(os.getenv("LOCAL_OBJECT_STORE_DIR", "app/data/object_storage")) / "uploads"

EXT_TO_TYPE = {
    ".pdf": "pdf", ".docx": "docx", ".doc": "docx",
    ".xlsx": "xlsx", ".xls": "xlsx",
    ".png": "image", ".jpg": "image", ".jpeg": "image",
    ".gif": "image", ".bmp": "image", ".webp": "image",
}

EXT_TO_MIME = {
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".doc": "application/msword",
    ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    ".xls": "application/vnd.ms-excel",
    ".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
    ".gif": "image/gif", ".bmp": "image/bmp", ".webp": "image/webp",
}

EXT_TO_MEDIA = {
    ".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
    ".gif": "image/gif", ".bmp": "image/bmp", ".webp": "image/webp",
}


# ========== 视觉模型池（circuit-breaker failover）==========

class _VisionEndpoint:
    def __init__(self, name: str, base_url: str, token: str, model: str,
                 threshold: int, cooldown: int):
        self.name = name
        self.base_url = base_url.rstrip("/")
        self.token = token
        self.model = model
        self.threshold = threshold
        self.cooldown = cooldown
        self.failure_count = 0
        self.demoted_until: float = 0.0

    def is_available(self) -> bool:
        if not self.base_url or not self.token or not self.model:
            return False
        if self.demoted_until:
            if time.time() < self.demoted_until:
                return False
            # 冷却到期，自动恢复
            self.failure_count = 0
            self.demoted_until = 0.0
        return True

    def on_success(self):
        self.failure_count = 0
        self.demoted_until = 0.0

    def on_failure(self) -> bool:
        """返回 True 表示刚触发熔断"""
        self.failure_count += 1
        if self.failure_count >= self.threshold:
            self.demoted_until = time.time() + self.cooldown
            return True
        return False


class VisionModelPool:
    """进程级单例，管理 VL1/VL2/VL3 三个视觉端点的 failover。"""
    _instance: Optional["VisionModelPool"] = None
    _init_lock = threading.Lock()

    @classmethod
    def get(cls) -> "VisionModelPool":
        if cls._instance is None:
            with cls._init_lock:
                if cls._instance is None:
                    cls._instance = cls()
        return cls._instance

    def __init__(self):
        default_threshold = int(os.getenv("LLM_FAILOVER_THRESHOLD_VL", "3"))
        default_cooldown = int(os.getenv("LLM_FAILOVER_COOLDOWN_VL", "300"))
        self._endpoints: List[_VisionEndpoint] = []
        for suffix, name in [("_VL", "VL1"), ("_VL2", "VL2"), ("_VL3", "VL3")]:
            ep = _VisionEndpoint(
                name=name,
                base_url=os.getenv(f"ANTHROPIC_BASE_URL{suffix}", ""),
                token=os.getenv(f"ANTHROPIC_AUTH_TOKEN{suffix}", ""),
                model=os.getenv(f"ANTHROPIC_MODEL{suffix}", ""),
                threshold=int(os.getenv(f"LLM_FAILOVER_THRESHOLD{suffix}", str(default_threshold))),
                cooldown=int(os.getenv(f"LLM_FAILOVER_COOLDOWN{suffix}", str(default_cooldown))),
            )
            self._endpoints.append(ep)
        configured = [ep.name for ep in self._endpoints if ep.base_url and ep.token and ep.model]
        logger.info(f"[VisionPool] Initialized: {configured}")

    def all_demoted(self) -> bool:
        return all(not ep.is_available() for ep in self._endpoints)

    async def call(self, file_path: str) -> Optional[str]:
        """依次尝试可用端点，成功返回文本，全部失败返回 None。"""
        import httpx

        ext = Path(file_path).suffix.lower()
        media_type = EXT_TO_MEDIA.get(ext, "image/png")
        with open(file_path, "rb") as f:
            image_b64 = base64.b64encode(f.read()).decode("utf-8")

        for ep in self._endpoints:
            if not ep.is_available():
                logger.info(f"[VisionPool] {ep.name} demoted, skipping")
                continue
            try:
                async with httpx.AsyncClient(timeout=60) as client:
                    resp = await client.post(
                        f"{ep.base_url}/chat/completions",
                        headers={"Authorization": f"Bearer {ep.token}", "Content-Type": "application/json"},
                        json={
                            "model": ep.model,
                            "messages": [{
                                "role": "user",
                                "content": [
                                    {"type": "text", "text": "请详细识别并描述这张图片中的所有文字内容和关键信息。如果包含表格，请以结构化方式输出。"},
                                    {"type": "image_url", "image_url": {"url": f"data:{media_type};base64,{image_b64}"}},
                                ],
                            }],
                            "max_tokens": 4096,
                        },
                    )
                text = resp.json().get("choices", [{}])[0].get("message", {}).get("content", "")
                if not text:
                    raise ValueError(f"empty response from {ep.model}")
                ep.on_success()
                logger.info(f"[VisionPool] {ep.name} ({ep.model}) OK: {len(text)} chars")
                return text
            except Exception as e:
                tripped = ep.on_failure()
                if tripped:
                    logger.warning(f"[VisionPool] {ep.name} tripped circuit-breaker (cooldown={ep.cooldown}s): {e}")
                else:
                    logger.warning(f"[VisionPool] {ep.name} failed ({ep.failure_count}/{ep.threshold}): {e}")

        logger.warning("[VisionPool] All vision endpoints failed or demoted")
        return None


# ========== 本地存储 ==========

def _persist_file(file_data: bytes, file_name: str) -> str:
    try:
        dest = UPLOAD_DIR / "1" / file_name
        dest.parent.mkdir(parents=True, exist_ok=True)
        dest.write_bytes(file_data)
        base_url = (os.getenv("AGENT_PUBLIC_BASE_URL") or
                    f"http://{os.getenv('AGENT_EXTERNAL_HOST', '127.0.0.1')}:{os.getenv('AGENT_SERVICE_PORT', '8000')}")
        from urllib.parse import quote
        token = f"object_storage/uploads/1/{file_name}"
        url = f"{base_url.rstrip('/')}/api/files/download?path={quote(token, safe='/:_-.()')}"
        logger.info(f"[DocumentReader] Persisted: {url}")
        return url
    except Exception as e:
        logger.warning(f"[DocumentReader] Persist failed: {e}")
        return ""


# ========== 文件类型检测 ==========

def _detect_file_type(file_path: str) -> Optional[str]:
    ext = Path(file_path).suffix.lower()
    ft = EXT_TO_TYPE.get(ext)
    if ft:
        return ft
    try:
        with open(file_path, "rb") as f:
            h = f.read(16)
        if h[:8] == b'\x89PNG\r\n\x1a\n': return "image"
        if h[:3] == b'\xff\xd8\xff': return "image"
        if h[:4] == b'%PDF': return "pdf"
        if h[:4] == b'PK\x03\x04': return "docx"
        if h[:3] == b'GIF': return "image"
    except Exception:
        pass
    return None


# ========== 文件下载 ==========

async def _download_url(url: str) -> tuple:
    import httpx
    from urllib.parse import urlparse, unquote

    UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
    async with httpx.AsyncClient(timeout=60, follow_redirects=True) as client:
        resp = await client.get(url)
        resp.raise_for_status()

    # 文件名提取
    file_name = ""
    cd = resp.headers.get("content-disposition", "")
    if "filename=" in cd:
        file_name = cd.split("filename=")[-1].strip().strip('"').strip("'")
    if not file_name:
        path_name = unquote(Path(urlparse(url).path).name)
        if path_name and "." in path_name:
            file_name = path_name
    if not file_name:
        ct = resp.headers.get("content-type", "")
        ext = next((v for k, v in {
            "application/pdf": ".pdf", "image/png": ".png", "image/jpeg": ".jpg",
            "image/gif": ".gif", "image/webp": ".webp", "image/bmp": ".bmp",
        }.items() if k in ct), ".bin")
        file_name = f"download_{int(time.time())}{ext}"

    ext = Path(file_name).suffix.lower() or ".bin"
    save_path = UPLOAD_DIR / f"{int(time.time())}_{file_name}"
    save_path.parent.mkdir(parents=True, exist_ok=True)
    save_path.write_bytes(resp.content)
    logger.info(f"[DocumentReader] Downloaded: {save_path} ({len(resp.content)} bytes)")
    return str(save_path), file_name


async def _download_dingtalk(download_code: str, robot_code: str) -> tuple:
    import httpx

    client_id = os.getenv("DINGTALK_CLIENT_ID", "")
    client_secret = os.getenv("DINGTALK_CLIENT_SECRET", "")
    if not client_id or not client_secret:
        raise RuntimeError("DINGTALK_CLIENT_ID / DINGTALK_CLIENT_SECRET not configured")

    async with httpx.AsyncClient(timeout=10) as client:
        r = await client.post(
            "https://api.dingtalk.com/v1.0/oauth2/accessToken",
            json={"appKey": client_id, "appSecret": client_secret},
        )
        token = r.json().get("accessToken", "")
    if not token:
        raise RuntimeError("获取钉钉 access_token 失败")

    async with httpx.AsyncClient(timeout=30) as client:
        r = await client.post(
            "https://api.dingtalk.com/v1.0/robot/messageFiles/download",
            headers={"x-acs-dingtalk-access-token": token, "Content-Type": "application/json"},
            json={"downloadCode": download_code, "robotCode": robot_code},
        )
        data = r.json()
    dl_url = data.get("downloadUrl")
    if not dl_url:
        raise RuntimeError(f"钉钉 downloadCode 无效: {data}")
    return await _download_url(dl_url)


# ========== 格式解析器 ==========

async def _parse_pdf(file_path: str, max_chars: int) -> Dict[str, Any]:
    text_parts, tables, page_count, metadata = [], [], 0, {}
    try:
        import pdfplumber
        with pdfplumber.open(file_path) as pdf:
            page_count = len(pdf.pages)
            metadata = pdf.metadata or {}
            for i, page in enumerate(pdf.pages):
                text = page.extract_text()
                if text and text.strip():
                    text_parts.append(text.strip())
                for tbl in page.extract_tables():
                    if tbl and len(tbl) > 1:
                        tables.append({"location": f"Page {i+1}", "headers": [str(h or "") for h in tbl[0]], "rows": [[str(c or "") for c in r] for r in tbl[1:]][:50]})
                if len("\n".join(text_parts)) >= max_chars:
                    break
    except Exception as e:
        logger.warning(f"[DocumentReader] pdfplumber failed: {e}")
        try:
            import pypdf
            reader = pypdf.PdfReader(file_path)
            page_count = len(reader.pages)
            metadata = dict(reader.metadata) if reader.metadata else {}
            for page in reader.pages:
                text = page.extract_text()
                if text and text.strip():
                    text_parts.append(text.strip())
                if len("\n".join(text_parts)) >= max_chars:
                    break
        except Exception as e2:
            logger.error(f"[DocumentReader] pypdf also failed: {e2}")
    return {
        "file_type": "pdf", "file_size": os.path.getsize(file_path),
        "page_count": page_count, "text_content": "\n\n".join(text_parts)[:max_chars],
        "tables": tables[:20], "metadata": {k: str(v) for k, v in metadata.items() if v},
    }


async def _parse_docx(file_path: str, max_chars: int) -> Dict[str, Any]:
    from docx import Document
    doc = Document(file_path)
    text_parts, tables = [], []
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text.strip())
        if len("\n".join(text_parts)) >= max_chars:
            break
    for i, tbl in enumerate(doc.tables):
        headers, rows_data = [], []
        for j, row in enumerate(tbl.rows):
            cells = [c.text.strip() for c in row.cells]
            if j == 0:
                headers = cells
            else:
                rows_data.append(cells)
        if headers:
            tables.append({"location": f"Table {i+1}", "headers": headers, "rows": rows_data[:50]})
    return {
        "file_type": "docx", "file_size": os.path.getsize(file_path),
        "text_content": "\n".join(text_parts)[:max_chars], "tables": tables[:20],
        "metadata": {"paragraph_count": len(doc.paragraphs), "table_count": len(doc.tables)},
    }


async def _parse_xlsx(file_path: str, max_chars: int) -> Dict[str, Any]:
    import openpyxl
    wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
    sheets_info, text_parts = [], []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            sheets_info.append({"sheet": sheet_name, "rows": 0, "cols": 0})
            continue
        headers = [str(h or "") for h in rows[0]]
        data_rows = [[str(c or "") for c in r] for r in rows[1:]]
        sheets_info.append({"sheet": sheet_name, "rows": len(data_rows), "cols": len(headers), "headers": headers, "sample_rows": data_rows[:30]})
        text_parts.append(f"--- Sheet: {sheet_name} ---\n" + " | ".join(headers) + "\n" + "\n".join(" | ".join(r) for r in data_rows[:30]))
    wb.close()
    return {
        "file_type": "xlsx", "file_size": os.path.getsize(file_path),
        "text_content": "\n\n".join(text_parts)[:max_chars],
        "sheets": sheets_info, "tables": sheets_info,
        "metadata": {"sheet_count": len(sheets_info)},
    }


async def _parse_image(file_path: str, max_chars: int) -> Dict[str, Any]:
    """视觉模型优先（VL1→VL2→VL3），全部熔断时降级到 pytesseract OCR。"""
    pool = VisionModelPool.get()

    if not pool.all_demoted():
        vl_text = await pool.call(file_path)
        if vl_text:
            return {
                "file_type": "image", "file_size": os.path.getsize(file_path),
                "text_content": vl_text[:max_chars],
                "recognition_method": "vision_model", "recognition_model": "vl_pool",
            }
    else:
        logger.info("[DocumentReader] All VL endpoints demoted, using OCR directly")

    ocr_text = ""
    try:
        import pytesseract
        from PIL import Image
        ocr_text = pytesseract.image_to_string(Image.open(file_path), lang="chi_sim+eng").strip()
        logger.info(f"[DocumentReader] OCR fallback: {len(ocr_text)} chars")
    except ImportError:
        logger.warning("[DocumentReader] pytesseract not installed")
    except Exception as e:
        logger.warning(f"[DocumentReader] pytesseract failed: {e}")

    return {
        "file_type": "image", "file_size": os.path.getsize(file_path),
        "text_content": (ocr_text or "图片识别失败")[:max_chars],
        "recognition_method": "ocr", "recognition_model": "pytesseract",
    }


# ========== 主入口 ==========

async def _run(params: Dict[str, Any]) -> Dict[str, Any]:
    local_path = None
    original_file_name = None

    try:
        if params.get("file_path"):
            local_path = params["file_path"]
            original_file_name = os.path.basename(local_path)
        elif params.get("file_url"):
            local_path, original_file_name = await _download_url(params["file_url"])
        elif params.get("dingtalk_download_code"):
            local_path, original_file_name = await _download_dingtalk(
                params["dingtalk_download_code"],
                params.get("dingtalk_robot_code", ""),
            )
        else:
            return {"error": "缺少文件来源参数（file_path / file_url / dingtalk_download_code）",
                    "for_llm": {"error": "未提供文件来源"}}

        if not local_path or not os.path.exists(local_path):
            return {"error": f"文件获取失败: {local_path}", "for_llm": {"error": "文件获取失败"}}

        if not original_file_name:
            original_file_name = os.path.basename(local_path)

        # 持久化
        with open(local_path, "rb") as f:
            file_data = f.read()
        download_url = _persist_file(file_data, original_file_name)

        # 检测类型
        file_type = params.get("file_type", "auto")
        if file_type == "auto":
            file_type = _detect_file_type(local_path)
        if not file_type:
            return {"error": f"无法识别文件类型: {local_path}", "for_llm": {"error": "无法识别文件类型"}}

        max_chars = int(params.get("max_chars", 10000))
        extract_mode = params.get("extract_mode", "full")

        parsers = {"pdf": _parse_pdf, "docx": _parse_docx, "xlsx": _parse_xlsx, "image": _parse_image}
        parser = parsers.get(file_type)
        if not parser:
            return {"error": f"不支持的文件类型: {file_type}", "for_llm": {"error": f"不支持: {file_type}"}}

        result = await parser(local_path, max_chars)
        result["file_name"] = original_file_name
        result["file_path"] = local_path
        result["download_url"] = download_url

        if extract_mode == "summary" and result.get("text_content"):
            result["text_content"] = result["text_content"][:5000]

        logger.info(f"[DocumentReader] Done: {original_file_name} ({file_type}), text={len(result.get('text_content', ''))} chars")

        return {
            "for_llm": {
                "file_name": original_file_name,
                "file_type": file_type,
                "text_content": result.get("text_content", ""),
                "download_url": download_url,
            },
            **result,
        }

    except Exception as e:
        import traceback
        traceback.print_exc()
        logger.error(f"[DocumentReader] Failed: {e}")
        return {"error": str(e), "for_llm": {"error": f"文档处理失败: {e}"}}


def main(params: Dict[str, Any]) -> Dict[str, Any]:
    return asyncio.run(_run(params))


if __name__ == "__main__":
    import argparse

    p: Dict[str, Any] = {}
    if not sys.stdin.isatty():
        try:
            raw = sys.stdin.read().strip()
            if raw:
                p = json.loads(raw)
        except Exception:
            pass

    parser = argparse.ArgumentParser()
    parser.add_argument("--file-path", dest="file_path")
    parser.add_argument("--file-url", dest="file_url")
    parser.add_argument("--dingtalk-download-code", dest="dingtalk_download_code")
    parser.add_argument("--dingtalk-robot-code", dest="dingtalk_robot_code")
    args = parser.parse_args()
    for k, v in vars(args).items():
        if v is not None:
            p[k] = v

    result = main(p)
    print(json.dumps(result, ensure_ascii=False, default=str, indent=2))
